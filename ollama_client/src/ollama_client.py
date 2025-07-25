import streamlit as st
from PIL import Image
import io
import base64
import pandas as pd
from PyPDF2 import PdfReader
import docx
import openpyxl
import requests

# 개인정보 정의/예시
PERSONAL_INFO_CONTEXT = """
개인정보란 이름, 주민등록번호, 연락처, 이메일 등 개인을 식별할 수 있는 정보를 의미합니다.
개인정보에는 주소, 전화번호, 계좌번호, 신용카드번호, 생년월일, 성별 등이 포함될 수 있습니다.
개인정보 보호법에 따라 개인을 식별할 수 있는 모든 정보는 보호 대상입니다.
개인정보의 예시: 김철수, 010-1234-5678, kim@email.com, 서울시 강남구, 123-45-67890
개인정보가 아닌 것: 일반적인 직업명, 나이대, 지역명(시/도 단위), 성별 등
"""

OLLAMA_API_URL = "http://localhost:11434/api/generate"
MODEL_NAME = "qwen2.5vl:7b"
#MODEL_NAME = "qwen3:8b" 

def encode_image_to_base64(image):
    buffered = io.BytesIO()
    image.save(buffered, format="PNG")
    return base64.b64encode(buffered.getvalue()).decode()

def qwen_ocr(image):
    """qwen2.5vl:7b로 이미지 내 문자 추출"""
    prompt = "이 이미지에 포함된 모든 문자를 정확하게 추출해서 반환해줘. (텍스트만 반환)"
    payload = {
        "model": MODEL_NAME,
        "prompt": prompt,
        "images": [encode_image_to_base64(image)],
        "stream": False
    }
    try:
        response = requests.post(OLLAMA_API_URL, json=payload)
        response.raise_for_status()
        return response.json()["response"]
    except Exception as e:
        return f"OCR Error: {str(e)}"

def extract_images_from_docx(file):
    doc = docx.Document(file)
    images = []
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            img_bytes = rel.target_part.blob
            images.append(Image.open(io.BytesIO(img_bytes)))
    return images

def extract_images_from_excel(file):
    wb = openpyxl.load_workbook(file)
    images = []
    for ws in wb.worksheets:
        for img in getattr(ws, '_images', []):
            if hasattr(img, 'ref'):
                img_bytes = img._data()
                images.append(Image.open(io.BytesIO(img_bytes)))
    return images

def extract_text_from_pdf(file):
    file.seek(0)
    reader = PdfReader(file)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text

def extract_text_from_docx(file):
    doc = docx.Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_text_from_excel(file):
    df = pd.read_excel(file)
    return df.to_string(index=False)

def generate_response(prompt, image=None):
    payload = {
        "model": MODEL_NAME,
        "prompt": prompt,
        "stream": False
    }
    if image:
        # 이미지 크기 제한 (예: 1024x1024)
        image = image.copy()
        image.thumbnail((1024, 1024))
        payload["images"] = [encode_image_to_base64(image)]
    try:
        response = requests.post(OLLAMA_API_URL, json=payload, timeout=120)  # 타임아웃 추가
        response.raise_for_status()
        return response.json()["response"]
    except requests.Timeout:
        return "Ollama 서버 응답이 너무 오래 걸립니다(타임아웃). 서버 상태를 확인하세요."
    except Exception as e:
        return f"Error: {str(e)}"

st.set_page_config(page_title="On-Device 개인정보 가드레일 시스템", layout="wide")
st.markdown("""
<style>
    .block-container { padding-top: 2rem; }
    .stChatMessage { background: #f7f7f8; border-radius: 1.2em; margin-bottom: 1.2em; box-shadow: 0 2px 8px #0001; }
    .stChatMessage.user { background: #e6f0ff; }
    .stChatMessage.assistant { background: #f7f7f8; }
    .stChatInputContainer { position: fixed; bottom: 0; left: 0; right: 0; background: #fff; padding: 1em 0.5em 0.5em 0.5em; box-shadow: 0 -2px 8px #0001; z-index: 100; }
</style>
""", unsafe_allow_html=True)

st.title("🔒 On-Device 개인정보 가드레일 시스템 ")

if "messages" not in st.session_state:
    st.session_state.messages = []

# 메시지 렌더링
for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        if msg["type"] == "text":
            st.markdown(msg["content"])
        elif msg["type"] == "image":
            st.image(msg["content"], caption="업로드된 이미지", use_column_width=True)
        elif msg["type"] == "file_text":
            st.text_area(msg.get("label", "파일 내용"), msg["content"], height=200, key=msg.get("key", None), disabled=True)
        elif msg["type"] == "result":
            if "포함됨" in msg["content"]:
                st.markdown("""
                <div style="background-color: #ffebee; border-left: 4px solid #f44336; padding: 16px; margin: 8px 0; border-radius: 4px;">
                    <h4 style="color: #c62828; margin: 0;">🚨 개인정보가 포함되어 있습니다!</h4>
                    <p style="color: #d32f2f; margin: 8px 0 0 0; font-weight: 500;">민감한 개인정보가 감지되었습니다. 주의가 필요합니다.</p>
                </div>
                """, unsafe_allow_html=True)
            elif "포함되지 않음" in msg["content"]:
                st.markdown("""
                <div style="background-color: #e8f5e8; border-left: 4px solid #4caf50; padding: 16px; margin: 8px 0; border-radius: 4px;">
                    <h4 style="color: #2e7d32; margin: 0;">✅ 개인정보가 포함되어 있지 않습니다!</h4>
                    <p style="color: #388e3c; margin: 8px 0 0 0; font-weight: 500;">안전합니다. 개인정보가 감지되지 않았습니다.</p>
                </div>
                """, unsafe_allow_html=True)
            else:
                st.markdown("""
                <div style="background-color: #fff3e0; border-left: 4px solid #ff9800; padding: 16px; margin: 8px 0; border-radius: 4px;">
                    <h4 style="color: #e65100; margin: 0;">⚠️ 판별 결과를 확인할 수 없습니다.</h4>
                    <p style="color: #f57c00; margin: 8px 0 0 0; font-weight: 500;">결과를 다시 확인해주세요.</p>
                </div>
                """, unsafe_allow_html=True)
            st.markdown("**📋 상세 분석 결과:**")
            st.markdown(f"""
            <div style="background-color: #f8f9fa; border: 1px solid #dee2e6; border-radius: 6px; padding: 12px; margin: 8px 0;">
                <code style="color: #495057; font-size: 14px;">{msg["content"]}</code>
            </div>
            """, unsafe_allow_html=True)

# 하단 입력창/업로드
with st.container():
    st.markdown("<div class='stChatInputContainer'>", unsafe_allow_html=True)
    col_upload, col_input, col_btn = st.columns([1, 6, 1])
    with col_upload:
        uploaded_file = st.file_uploader("", type=["jpg", "jpeg", "png", "pdf", "docx", "xlsx", "txt"], label_visibility="collapsed", key="file_uploader")
    with col_input:
        user_text = st.chat_input("메시지 입력 또는 파일 업로드...")
    with col_btn:
        analyze_btn = st.button("🚀 판별", use_container_width=True)
    st.markdown("</div>", unsafe_allow_html=True)

# 파일 업로드/텍스트 입력 처리
if uploaded_file is not None:
    filetype = uploaded_file.type
    filename = uploaded_file.name.lower()
    image = None
    ocr_texts = []
    user_text_from_file = ""
    # 이미지 파일
    if filetype.startswith("image/"):
        image = Image.open(uploaded_file)
        st.session_state.messages.append({"role": "user", "type": "image", "content": image})
    elif filename.endswith(".pdf"):
        try:
            from pdf2image import convert_from_bytes
            images = convert_from_bytes(uploaded_file.read())
            for idx, img in enumerate(images):
                st.session_state.messages.append({"role": "user", "type": "image", "content": img})
                ocr_result = qwen_ocr(img)
                ocr_texts.append(ocr_result)
            uploaded_file.seek(0)
        except ImportError:
            st.session_state.messages.append({"role": "assistant", "type": "text", "content": "pdf2image/poppler 미설치로 PDF 내 이미지 OCR은 생략됩니다."})
        user_text_from_file = extract_text_from_pdf(uploaded_file)
        st.session_state.messages.append({"role": "user", "type": "file_text", "content": user_text_from_file, "label": "PDF에서 추출된 텍스트", "key": "pdf_text"})
    elif filename.endswith(".docx"):
        images = extract_images_from_docx(uploaded_file)
        for idx, img in enumerate(images):
            st.session_state.messages.append({"role": "user", "type": "image", "content": img})
            ocr_result = qwen_ocr(img)
            ocr_texts.append(ocr_result)
        user_text_from_file = extract_text_from_docx(uploaded_file)
        st.session_state.messages.append({"role": "user", "type": "file_text", "content": user_text_from_file, "label": "DOCX에서 추출된 텍스트", "key": "docx_text"})
    elif filename.endswith(".xlsx"):
        images = extract_images_from_excel(uploaded_file)
        for idx, img in enumerate(images):
            st.session_state.messages.append({"role": "user", "type": "image", "content": img})
            ocr_result = qwen_ocr(img)
            ocr_texts.append(ocr_result)
        user_text_from_file = extract_text_from_excel(uploaded_file)
        st.session_state.messages.append({"role": "user", "type": "file_text", "content": user_text_from_file, "label": "Excel에서 추출된 텍스트", "key": "excel_text"})
    elif filename.endswith(".txt"):
        user_text_from_file = uploaded_file.read().decode("utf-8")
        st.session_state.messages.append({"role": "user", "type": "file_text", "content": user_text_from_file, "label": "TXT 파일 내용", "key": "txt_text"})
    else:
        st.session_state.messages.append({"role": "assistant", "type": "text", "content": "지원하지 않는 파일 형식입니다."})
    # 판별 버튼 클릭 시 자동 분석
    if analyze_btn or uploaded_file:
        if (image or (filename.endswith(".txt"))) and not ocr_texts:
            prompt = f"""다음은 개인정보의 정의와 예시입니다:\n{PERSONAL_INFO_CONTEXT}\n\n아래 입력(텍스트 또는 이미지)에 개인정보가 포함되어 있으면 '포함됨', 포함되어 있지 않으면 '포함되지 않음'이라고만 답하세요.\n\n입력:\n{user_text_from_file if user_text_from_file else '[이미지 첨부]'}\n"""
            with st.spinner("분석 중..."):
                result = generate_response(prompt, image)
        elif user_text_from_file or ocr_texts:
            all_text = user_text_from_file + "\n" + "\n".join(ocr_texts)
            prompt = f"""다음은 개인정보의 정의와 예시입니다:\n{PERSONAL_INFO_CONTEXT}\n\n아래 입력(텍스트 또는 이미지)에 개인정보가 포함되어 있으면 '포함됨', 포함되어 있지 않으면 '포함되지 않음'이라고만 답하세요.\n\n입력:\n{all_text if all_text.strip() else '[이미지 첨부]'}\n"""
            with st.spinner("분석 중..."):
                result = generate_response(prompt)
        else:
            result = None
        if result:
            st.session_state.messages.append({"role": "assistant", "type": "result", "content": result})
        st.rerun()

# 텍스트 입력 처리
if user_text:
    st.session_state.messages.append({"role": "user", "type": "text", "content": user_text})
    prompt = f"""다음은 개인정보의 정의와 예시입니다:\n{PERSONAL_INFO_CONTEXT}\n\n아래 입력(텍스트 또는 이미지)에 개인정보가 포함되어 있으면 '포함됨', 포함되어 있지 않으면 '포함되지 않음'이라고만 답하세요.\n\n입력:\n{user_text}\n"""
    with st.spinner("분석 중..."):
        result = generate_response(prompt)
    if result:
        st.session_state.messages.append({"role": "assistant", "type": "result", "content": result})
    st.rerun()

# 안내
with st.expander("ℹ️ 사용법/지원 포맷", expanded=False):
    st.markdown("""
    - 이미지, PDF, DOCX, Excel, 텍스트 모두 지원
    - Local sLLM 모델이 실행 중이어야 함
    - 판별 기준: 이름, 연락처, 이메일, 주소, 계좌번호 등
    """) 